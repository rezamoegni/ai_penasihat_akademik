# -*- coding: utf-8 -*-
"""
Penasihat Akademik SMA
--------------------------------
Aplikasi chatbot analisis profil akademik berbasis Google Gemini, LangChain, dan Streamlit.
Fungsi: Memberi rekomendasi bidang/jurusan kuliah berdasarkan kekuatan mata pelajaran, minat,
gaya belajar, dan preferensi siswa. Mendukung upload dokumen opsional (rapor/sertifikat)
untuk memperkaya konteks (RAG).

Cara menjalankan:
1) pip install -U streamlit google-generativeai langchain langchain-google-genai langchain-community chromadb PyPDF2 python-docx
2) streamlit run ai_penasihat_akademik.py
3) Masukkan Google AI API Key di sidebar.

Catatan:
- Aplikasi ini TIDAK mengambil data internet. Rekomendasi adalah kombinasi pemetaan berbasis aturan
  + reasoning dari model Gemini dengan konteks profil siswa.
"""

import os
from datetime import datetime

import streamlit as st

# Paket pemrosesan dokumen (opsional untuk upload rapor/sertifikat)
import PyPDF2
import docx

# Google Gemini (langsung)
import google.generativeai as genai

# LangChain (LLM + Embeddings + VectorStore + Prompting)
try:
    from langchain_google_genai import (
        ChatGoogleGenerativeAI,
        GoogleGenerativeAIEmbeddings,
    )
    from langchain_community.vectorstores import Chroma
    from langchain.text_splitter import RecursiveCharacterTextSplitter
    from langchain_core.prompts import ChatPromptTemplate
    from langchain_core.output_parsers import StrOutputParser
    from langchain_core.runnables import RunnablePassthrough
    from langchain_core.documents import Document
except ImportError:
    st.error(
        "❗ Paket belum lengkap. Jalankan:\n\n"
        "pip install -U langchain langchain-google-genai langchain-community chromadb"
    )
    st.stop()


# --------------------------------------------------------------------------------------
# KONFIGURASI HALAMAN & TEMA
# --------------------------------------------------------------------------------------
st.set_page_config(
    page_title="Penasihat Akademik SMA",
    page_icon="🎓",
    layout="centered",
)

# CSS sederhana: nuansa biru, bersih, ramah remaja
st.markdown(
    """
    <style>
        :root {
            --biru-utama: #1e88e5;
            --biru-muda: #e8f2ff;
            --biru-grad-1: #e6f0ff;
            --biru-grad-2: #ffffff;
            --teks-gelap: #0f172a;
            --teks-sedang: #334155;
        }
        .stApp {
            background: linear-gradient(180deg, var(--biru-grad-1), var(--biru-grad-2));
        }
        .judul-utama {
            text-align: center;
            margin-top: 0.5rem;
            margin-bottom: 0.75rem;
            color: var(--teks-gelap);
            font-weight: 800;
            letter-spacing: .2px;
        }
        .subjudul {
            text-align: center;
            color: var(--teks-sedang);
            margin-bottom: 1.25rem;
        }
        .kartu {
            background: white;
            border: 1px solid #e6eefc;
            border-radius: 14px;
            padding: 18px 18px 12px 18px;
            box-shadow: 0 8px 24px rgba(30,136,229,.05);
        }
        .chip {
            display: inline-block;
            padding: 4px 10px;
            margin: 0 6px 6px 0;
            background: var(--biru-muda);
            color: var(--biru-utama);
            border-radius: 999px;
            font-size: 0.83rem;
        }
        .stButton>button {
            background: linear-gradient(135deg, #1e88e5 0%, #1565c0 100%);
            color: white;
            border: 0;
            padding: 0.6rem 1rem;
            border-radius: 12px;
            font-weight: 600;
        }
        .stButton>button:hover {
            filter: brightness(1.05);
        }
        .tips {
            color: var(--teks-sedang);
            font-size: .92rem;
        }
        .label-biru {
            color: var(--biru-utama);
            font-weight: 700;
        }
    </style>
    """,
    unsafe_allow_html=True,
)

st.markdown('<h1 class="judul-utama">🎓 Penasihat Akademik SMA</h1>', unsafe_allow_html=True)
st.markdown(
    '<div class="subjudul">Bantu kamu menemukan <b>bidang kuliah</b> yang selaras dengan kekuatan dan minatmu.</div>',
    unsafe_allow_html=True,
)


# --------------------------------------------------------------------------------------
# STATE (Bahasa Indonesia)
# --------------------------------------------------------------------------------------
if "pesan" not in st.session_state:
    st.session_state.pesan = []
if "memproses" not in st.session_state:
    st.session_state.memproses = False
if "tampilkan_tindakan_cepat" not in st.session_state:
    st.session_state.tampilkan_tindakan_cepat = True
if "analisis_tunda" not in st.session_state:
    st.session_state.analisis_tunda = False
if "konten_dokumen" not in st.session_state:
    st.session_state.konten_dokumen = None
if "rag_rantai" not in st.session_state:
    st.session_state.rag_rantai = None
if "retriever" not in st.session_state:
    st.session_state.retriever = None
if "ringkasan_profil" not in st.session_state:
    st.session_state.ringkasan_profil = None
if "rekomendasi_awal" not in st.session_state:
    st.session_state.rekomendasi_awal = None


# --------------------------------------------------------------------------------------
# SIDEBAR: API KEY & AKSI
# --------------------------------------------------------------------------------------
with st.sidebar:
    st.header("Pengaturan")
    google_api_key = st.text_input("Google AI API Key", type="password", help="Masukkan API key kamu")
    if google_api_key:
        st.success("✅ API Key Tersambung")
    else:
        st.warning("⚠️ API Key Diperlukan")

    st.divider()
    if st.button("🧹 Bersihkan Obrolan"):
        st.session_state.pesan = []
        st.session_state.konten_dokumen = None
        st.session_state.rag_rantai = None
        st.session_state.retriever = None
        st.session_state.analisis_tunda = False
        st.session_state.memproses = False
        st.session_state.tampilkan_tindakan_cepat = True
        st.session_state.ringkasan_profil = None
        st.session_state.rekomendasi_awal = None
        st.success("Obrolan dibersihkan.")

    st.divider()
    st.subheader("Status Profil")
    if st.session_state.ringkasan_profil:
        st.success("📘 Profil siswa sudah terisi")
        st.caption(st.session_state.ringkasan_profil[:250] + ("..." if len(st.session_state.ringkasan_profil) > 250 else ""))
    else:
        st.info("Isi formulir profil di halaman utama untuk memulai.")

# Wajib ada API key
if not google_api_key:
    st.info("Masukkan Google AI API Key di sidebar untuk mulai menggunakan aplikasi.")
    st.stop()

# Inisialisasi Gemini
try:
    genai.configure(api_key=google_api_key)
    gemini_model = genai.GenerativeModel("gemini-2.5-flash")
except Exception as e:
    st.error(f"Gagal menginisialisasi Gemini: {e}")
    st.stop()


# --------------------------------------------------------------------------------------
# INISIALISASI LANGCHAIN (cache)
# --------------------------------------------------------------------------------------
@st.cache_resource
def inisialisasi_langchain(api_key: str):
    try:
        chat_model = ChatGoogleGenerativeAI(
            google_api_key=api_key,
            model="gemini-2.5-flash",
            temperature=0,
        )
        embeddings = GoogleGenerativeAIEmbeddings(
            google_api_key=api_key,
            model="models/gemini-embedding-exp-03-07",
        )
        return chat_model, embeddings
    except Exception as e:
        st.error(f"Gagal inisialisasi LangChain: {e}")
        return None, None


chat_model, embeddings = inisialisasi_langchain(google_api_key)


# --------------------------------------------------------------------------------------
# UTIL: Ekstraksi teks dokumen (opsional)
# --------------------------------------------------------------------------------------
def ekstrak_teks_pdf(file):
    try:
        reader = PyPDF2.PdfReader(file)
        teks = ""
        for halaman in reader.pages:
            teks += halaman.extract_text()
        return teks
    except Exception as e:
        return f"Error membaca PDF: {str(e)}"


def ekstrak_teks_docx(file):
    try:
        d = docx.Document(file)
        teks = ""
        for p in d.paragraphs:
            teks += p.text + "\n"
        return teks
    except Exception as e:
        return f"Error membaca DOCX: {str(e)}"


def ekstrak_teks_txt(file):
    try:
        return str(file.read(), "utf-8")
    except Exception as e:
        return f"Error membaca TXT: {str(e)}"


# --------------------------------------------------------------------------------------
# PROFIL INPUT (Form) — Fokus: kekuatan mata pelajaran + minat + preferensi
# --------------------------------------------------------------------------------------
st.markdown('<div class="kartu">', unsafe_allow_html=True)
st.subheader("🧭 Profil Akademik Kamu")

with st.form("form_profil"):
    col_kiri, col_kanan = st.columns(2)
    with col_kiri:
        nama = st.text_input("Nama (opsional)")
        tingkat = st.selectbox("Kelas", ["X", "XI", "XII"])
        gaya_belajar = st.multiselect(
            "Gaya belajar yang paling cocok",
            ["Visual", "Auditori", "Kinestetik", "Kolaboratif", "Mandiri"],
        )
    with col_kanan:
        minat_bidang = st.multiselect(
            "Bidang minat (pilih yang paling menarik)",
            ["Sains", "Teknologi", "Kesehatan", "Bisnis/Manajemen", "Sosial/Humaniora",
             "Hukum/Pemerintahan", "Seni/Desain", "Lingkungan", "Komunikasi/Media"],
        )
        toleransi_matematika = st.select_slider(
            "Kenyamanan dengan Matematika",
            options=["Rendah", "Sedang", "Tinggi"],
            value="Sedang",
        )

    st.markdown("### ⭐ Kekuatan Mata Pelajaran (beri skor 0–10)")
    col1, col2, col3 = st.columns(3)
    with col1:
        mtk = st.slider("Matematika", 0, 10, 5)
        fis = st.slider("Fisika", 0, 10, 5)
        kim = st.slider("Kimia", 0, 10, 5)
        bio = st.slider("Biologi", 0, 10, 5)
    with col2:
        tik = st.slider("Informatika/TIK", 0, 10, 5)
        eko = st.slider("Ekonomi", 0, 10, 5)
        akn = st.slider("Akuntansi", 0, 10, 5)
        geo = st.slider("Geografi", 0, 10, 5)
    with col3:
        sos = st.slider("Sosiologi", 0, 10, 5)
        sej = st.slider("Sejarah", 0, 10, 5)
        ind = st.slider("Bahasa Indonesia", 0, 10, 5)
        eng = st.slider("Bahasa Inggris", 0, 10, 5)

    st.markdown("### 📎 Lampirkan Rapor/Sertifikat (opsional)")
    unggahan = st.file_uploader(
        "Format didukung: PDF, DOCX, TXT (opsional)",
        type=["pdf", "docx", "txt"],
        help="Jika diunggah, kontennya akan dipakai untuk memperkaya analisis.",
    )

    kirim = st.form_submit_button("🔍 Analisis Rekomendasi")

st.markdown('</div>', unsafe_allow_html=True)


# --------------------------------------------------------------------------------------
# PEMETAAN BERBASIS ATURAN: Mata pelajaran → jurusan/bidang
# --------------------------------------------------------------------------------------
def skor_bidang_dari_map(nilai, preferensi, toleransi_mtk):
    """
    Menghitung skor awal berbagai bidang berdasarkan kekuatan mata pelajaran + preferensi.
    nilai: dict {mapel: skor 0-10}
    preferensi: list bidang
    toleransi_mtk: 'Rendah' | 'Sedang' | 'Tinggi'
    """

    # Bobot per bidang (sederhana & transparan)
    peta = {
        "Kedokteran": {"Biologi": 3, "Kimia": 2, "B. Inggris": 1},
        "Farmasi": {"Kimia": 3, "Biologi": 2, "Matematika": 1},
        "Keperawatan": {"Biologi": 2, "B. Indonesia": 1, "B. Inggris": 1},
        "Teknik Informatika / Ilmu Komputer": {"Matematika": 3, "TIK": 3, "Fisika": 1, "B. Inggris": 1},
        "Data Science / AI": {"Matematika": 3, "TIK": 3, "B. Inggris": 1},
        "Teknik Sipil": {"Matematika": 2, "Fisika": 2, "Geografi": 1},
        "Teknik Lingkungan / HSE": {"Kimia": 2, "Biologi": 1, "Geografi": 2, "Fisika": 1},
        "Teknik Industri": {"Matematika": 2, "Fisika": 2, "B. Inggris": 1},
        "Arsitektur": {"Matematika": 2, "Fisika": 1, "B. Indonesia": 1},
        "Perencanaan Wilayah & Kota": {"Geografi": 3, "Matematika": 1, "Sejarah": 1},
        "Manajemen/Marketing": {"Ekonomi": 2, "B. Indonesia": 1, "B. Inggris": 1},
        "Akuntansi/Keuangan": {"Akuntansi": 3, "Matematika": 2, "Ekonomi": 2},
        "Hukum": {"B. Indonesia": 2, "Sejarah": 2, "Sosiologi": 1},
        "Psikologi": {"Biologi": 1, "Sosiologi": 2, "Matematika": 1},
        "Ilmu Komunikasi": {"B. Indonesia": 2, "B. Inggris": 1, "Sejarah": 1},
        "HI (Hubungan Internasional)": {"B. Inggris": 2, "Sejarah": 2, "Sosiologi": 1},
        "Sastra/Filologi": {"B. Indonesia": 2, "B. Inggris": 2},
        "DKV/Desain": {"B. Indonesia": 1},  # kreatif → tidak dipetakan murni dari mapel; LLM akan menambah konteks
    }

    skor = {k: 0.0 for k in peta.keys()}
    # Konversi label mapel
    mapel_alias = {
        "Matematika": "Matematika",
        "Fisika": "Fisika",
        "Kimia": "Kimia",
        "Biologi": "Biologi",
        "TIK": "TIK",
        "Ekonomi": "Ekonomi",
        "Akuntansi": "Akuntansi",
        "Geografi": "Geografi",
        "Sosiologi": "Sosiologi",
        "Sejarah": "Sejarah",
        "B. Indonesia": "B. Indonesia",
        "B. Inggris": "B. Inggris",
    }

    for bidang, bobot_mapel in peta.items():
        total = 0.0
        for m, b in bobot_mapel.items():
            # ambil nilai mapel
            nm = mapel_alias.get(m, m)
            v = nilai.get(nm, 0)
            total += b * v
        skor[bidang] = total

    # Bonus preferensi minat
    preferensi_bonus = {
        "Kesehatan": ["Kedokteran", "Farmasi", "Keperawatan"],
        "Sains": ["Farmasi", "Psikologi", "Data Science / AI"],
        "Teknologi": ["Teknik Informatika / Ilmu Komputer", "Data Science / AI", "Teknik Industri"],
        "Bisnis/Manajemen": ["Manajemen/Marketing", "Akuntansi/Keuangan"],
        "Sosial/Humaniora": ["Hukum", "HI (Hubungan Internasional)", "Ilmu Komunikasi", "Sastra/Filologi", "Psikologi"],
        "Seni/Desain": ["DKV/Desain", "Arsitektur"],
        "Lingkungan": ["Teknik Lingkungan / HSE", "PWK (Perencanaan Wilayah & Kota)"],
        "Hukum/Pemerintahan": ["Hukum", "HI (Hubungan Internasional)"],
        "Komunikasi/Media": ["Ilmu Komunikasi", "DKV/Desain"],
    }

    for p in preferensi:
        for bidang in preferensi_bonus.get(p, []):
            if bidang in skor:
                skor[bidang] *= 1.08  # bonus 8%

    # Penalti/tuning untuk Matematika
    if toleransi_mtk == "Rendah":
        for bidang in ["Teknik Informatika / Ilmu Komputer", "Data Science / AI", "Teknik Sipil", "Teknik Industri", "Arsitektur", "Akuntansi/Keuangan"]:
            skor[bidang] *= 0.87  # kurangi 13%
    elif toleransi_mtk == "Tinggi":
        for bidang in ["Teknik Informatika / Ilmu Komputer", "Data Science / AI", "Teknik Sipil", "Teknik Industri", "Arsitektur", "Akuntansi/Keuangan"]:
            skor[bidang] *= 1.06  # tambah 6%

    return skor


def buat_ringkasan_profil(nama, tingkat, gaya_belajar, minat_bidang, toleransi_mtk, nilai_mapel):
    ringkas = []
    if nama:
        ringkas.append(f"Nama: {nama}")
    ringkas.append(f"Kelas: {tingkat}")
    ringkas.append(f"Gaya belajar: {', '.join(gaya_belajar) if gaya_belajar else '-'}")
    ringkas.append(f"Minat: {', '.join(minat_bidang) if minat_bidang else '-'}")
    ringkas.append(f"Kenyamanan Matematika: {toleransi_mtk}")
    ringkas.append(
        "Skor Mapel: " + ", ".join([f"{k} {v}/10" for k, v in nilai_mapel.items()])
    )
    return "\n".join(ringkas)


def buat_dokumen_langchain(teks: str, sumber: str):
    return [
        Document(
            page_content=teks,
            metadata={
                "source": sumber,
                "processed_at": datetime.now().isoformat(),
                "char_count": len(teks),
                "word_count": len(teks.split()),
            },
        )
    ]


def buat_rag_chain(dokumen, embeddings, chat_model):
    try:
        if not embeddings or not chat_model:
            st.error("❌ Embeddings atau Chat Model tidak tersedia.")
            return None, None

        splitter = RecursiveCharacterTextSplitter(
            chunk_size=2000,
            chunk_overlap=300,
            length_function=len,
            separators=["\n\n", "\n", " ", ""],
        )
        potongan = splitter.split_documents(dokumen)
        if not potongan:
            st.error("Tidak bisa memproses dokumen.")
            return None, None

        with st.spinner("Menyiapkan memori konteks..."):
            vs = Chroma.from_documents(documents=potongan, embedding=embeddings, persist_directory=None)
            retriever = vs.as_retriever(search_type="similarity", search_kwargs={"k": 8})

        def format_docs(docs):
            if not docs:
                return "Tidak ada konteks tambahan."
            return "\n".join([f"--- Konteks {i+1} ---\n{d.page_content}" for i, d in enumerate(docs)])

        prompt = ChatPromptTemplate.from_template(
            """Anda adalah penasihat akademik untuk siswa SMA di Indonesia.
Gunakan konteks profil berikut untuk menjawab secara spesifik, empatik, dan actionable.

Konteks Profil:
{context}

Pertanyaan Pengguna:
{question}

Instruksi:
- Jelaskan alasan rekomendasi (kaitkan dengan nilai mapel, minat, dan gaya belajar).
- Beri 3-5 rekomendasi jurusan/kelompok program studi, plus alternatif jika syarat tertentu kurang cocok.
- Sertakan contoh kegiatan ekstrakurikuler atau proyek yang bisa dicoba dalam 3-6 bulan.
- Jika pengguna minta perbandingan jurusan, paparkan perbedaan fokus, mata kuliah inti, dan prospek umum.
- Hindari klaim institusi tertentu; berikan saran generik (misal: "universitas dengan akreditasi baik untuk X").
Jawaban terstruktur dan mudah dibaca."""
        )

        rag = (
            {"context": retriever | format_docs, "question": RunnablePassthrough()}
            | prompt
            | chat_model
            | StrOutputParser()
        )
        return rag, retriever
    except Exception as e:
        st.error(f"Gagal membuat RAG chain: {e}")
        return None, None


# --------------------------------------------------------------------------------------
# PROSES KETIKA TOMBOL ANALISIS DIKLIK
# --------------------------------------------------------------------------------------
if kirim:
    st.session_state.memproses = True
    st.session_state.tampilkan_tindakan_cepat = False
    st.session_state.analisis_tunda = True

    # Susun nilai mapel
    nilai_mapel = {
        "Matematika": mtk,
        "Fisika": fis,
        "Kimia": kim,
        "Biologi": bio,
        "TIK": tik,
        "Ekonomi": eko,
        "Akuntansi": akn,
        "Geografi": geo,
        "Sosiologi": sos,
        "Sejarah": sej,
        "B. Indonesia": ind,
        "B. Inggris": eng,
    }

    ringkasan = buat_ringkasan_profil(
        nama=nama,
        tingkat=tingkat,
        gaya_belajar=gaya_belajar,
        minat_bidang=minat_bidang,
        toleransi_mtk=toleransi_matematika,
        nilai_mapel=nilai_mapel,
    )
    st.session_state.ringkasan_profil = ringkasan

    # Ekstrak teks unggahan (opsional)
    teks_lampiran = ""
    if unggahan is not None:
        ext = unggahan.name.split(".")[-1].lower()
        if ext == "pdf":
            teks_lampiran = ekstrak_teks_pdf(unggahan)
        elif ext == "docx":
            teks_lampiran = ekstrak_teks_docx(unggahan)
        elif ext == "txt":
            teks_lampiran = ekstrak_teks_txt(unggahan)
        if isinstance(teks_lampiran, str) and not teks_lampiran.startswith("Error"):
            st.session_state.konten_dokumen = teks_lampiran
        else:
            st.warning("Gagal mengekstrak teks lampiran. Analisis tetap dilanjutkan tanpa lampiran.")

    # Bangun dokumen RAG (profil + lampiran jika ada)
    doks = []
    doks += buat_dokumen_langchain(st.session_state.ringkasan_profil, "profil_siswa.txt")
    if st.session_state.konten_dokumen:
        doks += buat_dokumen_langchain(st.session_state.konten_dokumen, unggahan.name)

    # RAG
    rag, retr = buat_rag_chain(doks, embeddings, chat_model)
    st.session_state.rag_rantai = rag
    st.session_state.retriever = retr

    # Rekomendasi awal berbasis aturan
    skor = skor_bidang_dari_map(nilai_mapel, minat_bidang, toleransi_matematika)
    urut = sorted(skor.items(), key=lambda x: x[1], reverse=True)
    top5 = [b for b, _ in urut[:5]]
    st.session_state.rekomendasi_awal = top5

    # Prompt utama ke Gemini (gabungkan rule-based + profil)
    prompt_awal = f"""
Kamu adalah penasihat akademik untuk siswa SMA di Indonesia.
Berikut profil ringkas siswa:

{st.session_state.ringkasan_profil}

Hasil pemetaan awal (rule-based) memberi kandidat teratas:
- {top5[0] if len(top5)>0 else '-'}
- {top5[1] if len(top5)>1 else '-'}
- {top5[2] if len(top5)>2 else '-'}
- {top5[3] if len(top5)>3 else '-'}
- {top5[4] if len(top5)>4 else '-'}

Tolong:
1) Validasi & pertajam 3-5 rekomendasi bidang/jurusan (boleh menambah/menyusun ulang).
2) Jelaskan alasan (hubungkan dengan nilai mapel, minat, gaya belajar, dan toleransi matematika).
3) Beri alternatif jika siswa ingin jalur yang lebih/kurang intensif Matematika.
4) Buat rencana aksi 90 hari (materi yang diperdalam, proyek mini, lomba/ekskul).
5) Hindari menyebut universitas spesifik; gunakan saran generik.

Susun jawaban ringkas, terstruktur (heading + bullet), dan ramah siswa.
"""
    try:
        with st.spinner("🔎 Menganalisis profil & menyusun rekomendasi..."):
            jawaban = gemini_model.generate_content(prompt_awal)
            st.session_state.pesan.append({"role": "assistant", "content": jawaban.text})
            st.success("✅ Rekomendasi siap! Silakan lanjut bertanya lewat chat di bawah.")
    except Exception as e:
        st.error(f"Gagal membuat rekomendasi awal: {e}")
    finally:
        st.session_state.memproses = False
        st.session_state.analisis_tunda = False


# --------------------------------------------------------------------------------------
# TINDAKAN CEPAT (hanya tampil di awal)
# --------------------------------------------------------------------------------------
if st.session_state.tampilkan_tindakan_cepat and len(st.session_state.pesan) == 0 and not st.session_state.memproses:
    st.subheader("⚡ Tindakan Cepat")
    c1, c2, c3 = st.columns(3)
    with c1:
        if st.button("🎯 Cocoknya Ambil Jurusan Apa?"):
            st.session_state.pesan.append({"role": "user", "content": "Berdasarkan profil saya, jurusan kuliah apa yang paling cocok? Jelaskan alasannya."})
    with c2:
        if st.button("🔁 Alternatif Minim Matematika"):
            st.session_state.pesan.append({"role": "user", "content": "Kalau saya kurang nyaman dengan matematika, apa alternatif jurusan yang tetap relevan dengan minat saya?"})
    with c3:
        if st.button("🧪 Ekskul & Proyek 3 Bulan"):
            st.session_state.pesan.append({"role": "user", "content": "Rekomendasikan kegiatan ekstrakurikuler dan proyek 3 bulan untuk menguji minat saya."})


# --------------------------------------------------------------------------------------
# ANTARMUKA CHAT
# --------------------------------------------------------------------------------------
st.subheader("💬 Konsultasi dengan AI Penasihat")
for m in st.session_state.pesan:
    if m["role"] == "user":
        with st.chat_message("user"):
            st.write(m["content"])
    else:
        with st.chat_message("assistant"):
            st.write(m["content"])

# Indikator proses
if st.session_state.memproses:
    with st.chat_message("assistant"):
        st.markdown("🤖 **AI sedang menulis jawaban...**")

# Input chat
if st.session_state.memproses:
    st.chat_input("Sedang memproses...", disabled=True)
else:
    if pertanyaan := st.chat_input("Tanya apa saja soal jurusan & kuliah..."):
        # Tambahkan pertanyaan pengguna
        st.session_state.pesan.append({"role": "user", "content": pertanyaan})
        st.session_state.memproses = True

        # Jalankan balasan (pakai RAG jika tersedia; fallback pakai ringkasan profil)
        try:
            if st.session_state.rag_rantai:
                with st.spinner("🤖 Menganalisis konteks profil kamu..."):
                    balasan = st.session_state.rag_rantai.invoke(pertanyaan)
                st.session_state.pesan.append({"role": "assistant", "content": balasan})
            else:
                # Fallback: gunakan Gemini dengan sistem prompt + ringkasan profil
                sistem = """Anda penasihat akademik SMA. Jawab spesifik sesuai profil.
Hindari menyebut kampus tertentu; berikan saran generik."""
                full_prompt = f"{sistem}\n\nProfil:\n{st.session_state.ringkasan_profil or '-'}\n\nPertanyaan: {pertanyaan}\n\nJawaban:"
                balasan = gemini_model.generate_content(full_prompt).text
                st.session_state.pesan.append({"role": "assistant", "content": balasan})
        except Exception as e:
            st.error(f"Gagal membuat jawaban: {e}")
        finally:
            st.session_state.memproses = False
            st.rerun()


# --------------------------------------------------------------------------------------
# PANEL INFO & TIPS
# --------------------------------------------------------------------------------------
st.divider()
st.markdown("**Apa yang bisa kubantu?**")
st.markdown(
    """
- Menyusun _shortlist_ jurusan yang cocok dengan kekuatanmu
- Alternatif jalur kalau kamu ingin porsi Matematika lebih/kurang
- Rencana aksi 90 hari: materi yang diperdalam, proyek mini, dan ekskul
- Menjawab pertanyaan perbandingan jurusan (fokus, mata kuliah inti, prospek umum)
    """
)
st.caption("Penasihat Akademik SMA • Didukung oleh Google Gemini + LangChain")